"""
Document Processing Service for text extraction.

Supports PDF, DOCX, TXT, and images (with OCR).
"""

import logging
import os
from typing import Optional

from django.conf import settings

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Service for extracting text from various document formats.
    """
    
    def __init__(self):
        self.tesseract_cmd = getattr(settings, 'TESSERACT_CMD', '/usr/bin/tesseract')
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from a document based on its type.
        
        Args:
            file_path: Path to the document file
            file_type: Type of document (pdf, docx, txt, image)
        
        Returns:
            Extracted text content
        """
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extractors = {
            'pdf': self._extract_from_pdf,
            'docx': self._extract_from_docx,
            'txt': self._extract_from_txt,
            'image': self._extract_from_image,
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        try:
            text = extractor(file_path)
            return self._clean_text(text)
        except Exception as e:
            logger.exception(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file."""
        
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            
            # If no text extracted, try OCR
            if len(text.strip()) < 50:
                logger.info(f"PDF has little text, attempting OCR: {file_path}")
                text = self._ocr_pdf(file_path)
            
            return text
            
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF extraction")
    
    def _ocr_pdf(self, file_path: str) -> str:
        """Apply OCR to a PDF file."""
        
        try:
            from pdf2image import convert_from_path
            import pytesseract
            
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
            
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=200)
            
            text_parts = []
            for i, image in enumerate(images):
                logger.debug(f"OCR processing page {i + 1}")
                text = pytesseract.image_to_string(image, lang='fra+eng')
                text_parts.append(text)
            
            return "\n\n".join(text_parts)
            
        except ImportError as e:
            logger.error(f"OCR dependencies not available: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            return "\n\n".join(text_parts)
            
        except ImportError:
            raise ImportError("python-docx is required for DOCX extraction")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file."""
        
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        raise ValueError(f"Could not decode file with any of: {encodings}")
    
    def _extract_from_image(self, file_path: str) -> str:
        """Extract text from image using OCR."""
        
        try:
            import pytesseract
            from PIL import Image
            
            pytesseract.pytesseract.tesseract_cmd = self.tesseract_cmd
            
            image = Image.open(file_path)
            
            # Enhance image for better OCR
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            text = pytesseract.image_to_string(image, lang='fra+eng')
            return text
            
        except ImportError:
            raise ImportError("pytesseract and Pillow are required for image OCR")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        
        if not text:
            return ""
        
        # Replace multiple whitespaces with single space
        import re
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Replace multiple newlines with double newline
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)
        
        # Remove very short lines (likely noise)
        lines = text.split('\n')
        lines = [l for l in lines if len(l) > 2 or l == '']
        text = '\n'.join(lines)
        
        return text.strip()
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about a file."""
        
        import os
        
        stat = os.stat(file_path)
        _, ext = os.path.splitext(file_path)
        
        return {
            'path': file_path,
            'size': stat.st_size,
            'extension': ext.lower().lstrip('.'),
            'modified': stat.st_mtime,
        }


class TextChunker:
    """
    Utility for chunking large texts into smaller pieces.
    """
    
    @staticmethod
    def chunk_text(
        text: str,
        max_chunk_size: int = 4000,
        overlap: int = 200
    ) -> list:
        """
        Split text into overlapping chunks.
        
        Args:
            text: Text to chunk
            max_chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
        
        Returns:
            List of text chunks
        """
        
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence end
                for char in ['. ', '.\n', '! ', '? ']:
                    last_pos = text.rfind(char, start, end)
                    if last_pos > start + max_chunk_size // 2:
                        end = last_pos + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def estimate_tokens(text: str) -> int:
        """Estimate the number of tokens in a text (rough approximation)."""
        
        # Rough estimate: ~4 characters per token
        return len(text) // 4


# Example usage
if __name__ == '__main__':
    processor = DocumentProcessor()
    
    # Test with a sample text file
    test_content = """
    Ceci est un texte de test pour vérifier l'extraction.
    
    Il contient plusieurs paragraphes avec du contenu varié.
    
    Le machine learning est une discipline fascinante.
    """
    
    # Create test file
    test_path = '/tmp/test_doc.txt'
    with open(test_path, 'w') as f:
        f.write(test_content)
    
    extracted = processor.extract_text(test_path, 'txt')
    print(f"Extracted text:\n{extracted}")
    
    # Test chunking
    long_text = test_content * 100
    chunks = TextChunker.chunk_text(long_text, max_chunk_size=500)
    print(f"\nChunked into {len(chunks)} pieces")
